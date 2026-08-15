# -*- coding: utf-8 -*-
"""
결과보고서 Markdown → Word(.docx) 변환기 (공식 서식 표 구조 유지).
사용: python tools/build_report_docx.py [입력.md] [출력.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

KR_FONT = "맑은 고딕"


def _set_kr_font(run, size=None, bold=None):
    run.font.name = KR_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_runs(paragraph, text, size=10, base_bold=False):
    """**굵게**와 `코드` 인라인 서식을 처리해 run 추가."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_kr_font(run, size, True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_kr_font(run, size)
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        else:
            run = paragraph.add_run(token)
            _set_kr_font(run, size, base_bold)


def split_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_table_line(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def is_align_line(line):
    return bool(re.fullmatch(r"\|?[\s:|-]+\|?", line.strip()))


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    # 기본 본문 스타일
    style = doc.styles["Normal"]
    style.font.name = KR_FONT
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), KR_FONT)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 표
        if is_table_line(line) and i + 1 < n and is_align_line(lines[i + 1]):
            header = split_table_row(line)
            i += 2
            rows = []
            while i < n and is_table_line(lines[i]):
                rows.append(split_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = "Table Grid"
            for c, text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(cell.paragraphs[0], text, size=9.5, base_bold=True)
            for r, row in enumerate(rows, start=1):
                for c in range(len(header)):
                    text = row[c] if c < len(row) else ""
                    add_runs(table.rows[r].cells[c].paragraphs[0], text, size=9.5)
            doc.add_paragraph()
            continue

        # 인용 안내
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run_text = stripped.lstrip("> ").strip()
            add_runs(p, run_text, size=9)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue

        # 수평선
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # 제목
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:
                h = doc.add_heading(level=0)
                add_runs(h, text, size=16, base_bold=True)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                h = doc.add_heading(level=min(level - 1, 3))
                add_runs(h, text, size=13 if level == 2 else 11.5, base_bold=True)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            i += 1
            continue

        # 글머리 목록
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1
            continue

        # 번호 목록
        m = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        # 일반 문단
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"docx written: {docx_path}")


def docx_to_pdf(docx_path: Path, pdf_path: Path):
    """MS Word COM 자동화로 PDF 변환 (wdFormatPDF=17)."""
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
        doc.SaveAs2(str(pdf_path.resolve()), FileFormat=17)
        doc.Close(False)
    finally:
        word.Quit()
        pythoncom.CoUninitialize()
    print(f"pdf written: {pdf_path}")


if __name__ == "__main__":
    base = Path(__file__).resolve().parent.parent
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else base / "docs" / "결과보고서_2026_오픈소스개발자대회.md"
    docx = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    pdf = md.with_suffix(".pdf")
    md_to_docx(md, docx)
    docx_to_pdf(docx, pdf)
